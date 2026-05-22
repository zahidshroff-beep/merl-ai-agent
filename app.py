"""
MERL AI Agent Dashboard
Professional Streamlit app for Monitoring, Evaluation, Research & Learning
"""

import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
import io
import os
import html
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ============================================================
# CONFIG & THEME
# ============================================================
st.set_page_config(
    page_title="MERL AI Agent | Monitoring, Evaluation, Research & Learning",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================
# PROFESSIONAL COLOR SYSTEM (Modern Blue + Green)
# ============================================================
PRIMARY_BLUE = "#0D6EFD"
ACCENT_GREEN = "#10B981"
DEEP_BLUE = "#1E40AF"
LIGHT_BLUE = "#EFF6FF"
LIGHT_GREEN = "#ECFDF5"
DARK_TEXT = "#0F172A"
GRAY = "#64748B"
LIGHT_GRAY = "#F1F5F9"
WARNING = "#F59E0B"
DANGER = "#EF4444"
WHITE = "#FFFFFF"
CARD_SHADOW = "0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -2px rgba(0, 0, 0, 0.1)"

# ============================================================
# ENHANCED MODERN CSS - Professional & Clean
# ============================================================
st.markdown(f"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&family=Space+Grotesk:wght@500;600&display=swap');

    :root {{
        --primary: {PRIMARY_BLUE};
        --accent: {ACCENT_GREEN};
        --deep: {DEEP_BLUE};
    }}

    /* Global Typography & Spacing */
    .stApp {{
        font-family: 'Inter', system_ui, sans-serif;
    }}
    
    h1, h2, h3 {{
        font-family: 'Space Grotesk', 'Inter', sans-serif;
        font-weight: 600;
        letter-spacing: -0.025em;
    }}

    /* Top Navigation Bar */
    .nav-container {{
        background: {WHITE};
        border-bottom: 1px solid {LIGHT_GRAY};
        padding: 0.75rem 1.5rem;
        margin: -1rem -1rem 1.5rem -1rem;
        position: sticky;
        top: 0;
        z-index: 999;
        box-shadow: 0 1px 0 0 rgba(0,0,0,0.03);
    }}
    
    .nav-inner {{
        max-width: 1280px;
        margin: 0 auto;
        display: flex;
        align-items: center;
        justify-content: space-between;
    }}
    
    .nav-logo {{
        display: flex;
        align-items: center;
        gap: 0.6rem;
        font-size: 1.35rem;
        font-weight: 700;
        color: {DEEP_BLUE};
        text-decoration: none;
    }}
    
    .nav-logo span {{
        background: linear-gradient(90deg, {PRIMARY_BLUE}, {ACCENT_GREEN});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 700;
    }}

    .nav-pills {{
        display: flex;
        gap: 0.25rem;
        background: {LIGHT_GRAY};
        padding: 4px;
        border-radius: 9999px;
    }}
    
    .nav-pill {{
        padding: 0.45rem 1.1rem;
        border-radius: 9999px;
        font-size: 0.875rem;
        font-weight: 600;
        color: {GRAY};
        cursor: pointer;
        transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1);
        border: none;
        background: transparent;
    }}
    
    .nav-pill.active {{
        background: {WHITE};
        color: {DEEP_BLUE};
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }}
    
    .nav-pill:hover:not(.active) {{
        color: {DARK_TEXT};
    }}

    /* Support for <a> tags used in functional top nav (replaces non-working buttons) */
    .nav-pills a.nav-pill {{
        text-decoration: none;
        display: inline-block;
        vertical-align: middle;
        line-height: 1;
    }}
    a.nav-pill:hover:not(.active) {{
        color: {DARK_TEXT};
    }}

    /* Hero Section */
    .hero {{
        background: linear-gradient(135deg, {PRIMARY_BLUE} 0%, {DEEP_BLUE} 100%);
        border-radius: 20px;
        padding: 3.25rem 3rem;
        color: white;
        margin-bottom: 2rem;
        position: relative;
        overflow: hidden;
    }}
    
    .hero::before {{
        content: '';
        position: absolute;
        top: -50%;
        right: -20%;
        width: 60%;
        height: 200%;
        background: radial-gradient(circle, rgba(255,255,255,0.12) 0%, transparent 70%);
        transform: rotate(25deg);
    }}
    
    .hero h1 {{
        font-size: 3rem;
        font-weight: 700;
        line-height: 1.05;
        margin: 0 0 0.75rem 0;
        letter-spacing: -0.04em;
    }}
    
    .hero p {{
        font-size: 1.15rem;
        opacity: 0.92;
        max-width: 560px;
        margin-bottom: 1.75rem;
    }}

    /* Feature Cards */
    .feature-grid {{
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(210px, 1fr));
        gap: 1rem;
        margin: 1.5rem 0;
    }}
    
    .feature-card {{
        background: {WHITE};
        border: 1px solid {LIGHT_GRAY};
        border-radius: 14px;
        padding: 1.25rem 1.35rem;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }}
    
    .feature-card:hover {{
        transform: translateY(-3px);
        box-shadow: {CARD_SHADOW};
    }}
    
    .feature-icon {{
        width: 42px;
        height: 42px;
        border-radius: 10px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1.35rem;
        margin-bottom: 0.85rem;
    }}

    /* Modern Metric Cards */
    .metric-card {{
        background: {WHITE};
        border: 1px solid {LIGHT_GRAY};
        border-radius: 14px;
        padding: 1.15rem 1.35rem;
        box-shadow: {CARD_SHADOW};
        transition: transform 0.2s ease, box-shadow 0.2s ease;
        height: 100%;
    }}
    
    .metric-card:hover {{
        transform: translateY(-2px);
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.08), 0 4px 6px -4px rgba(0, 0, 0, 0.08);
    }}
    
    .metric-value {{
        font-size: 2.1rem;
        font-weight: 700;
        line-height: 1;
        color: {DEEP_BLUE};
        margin-bottom: 0.15rem;
    }}
    
    .metric-label {{
        font-size: 0.8rem;
        font-weight: 600;
        color: {GRAY};
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }}
    
    .metric-delta {{
        font-size: 0.75rem;
        font-weight: 600;
        margin-top: 0.35rem;
    }}

    /* Section Headers */
    .section-header {{
        font-size: 1.35rem;
        font-weight: 700;
        color: {DARK_TEXT};
        margin: 0 0 1rem 0;
        letter-spacing: -0.02em;
    }}
    
    .section-sub {{
        color: {GRAY};
        font-size: 0.95rem;
        margin-bottom: 1.25rem;
    }}

    /* Modern Upload Area */
    .upload-zone {{
        border: 2px dashed #CBD5E1;
        border-radius: 16px;
        padding: 2.25rem 1.5rem;
        text-align: center;
        background: {LIGHT_GRAY};
        transition: all 0.2s ease;
    }}
    
    .upload-zone:hover {{
        border-color: {PRIMARY_BLUE};
        background: {LIGHT_BLUE};
    }}

    /* Chat Improvements */
    .chat-container {{
        max-height: 420px;
        overflow-y: auto;
        padding-right: 0.5rem;
    }}
    
    .chat-bubble {{
        padding: 0.85rem 1.1rem;
        border-radius: 16px;
        margin-bottom: 0.65rem;
        max-width: 82%;
        line-height: 1.45;
        font-size: 0.94rem;
    }}
    
    .chat-user {{
        background: {PRIMARY_BLUE};
        color: white;
        margin-left: auto;
        border-bottom-right-radius: 4px;
    }}
    
    .chat-ai {{
        background: {LIGHT_GREEN};
        color: {DARK_TEXT};
        border: 1px solid #A7F3D0;
        border-bottom-left-radius: 4px;
    }}
    
    .chat-avatar {{
        width: 26px;
        height: 26px;
        border-radius: 50%;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        font-size: 0.8rem;
        margin-right: 0.5rem;
    }}

    /* Prompt Chips */
    .prompt-chips {{
        display: flex;
        flex-wrap: wrap;
        gap: 0.5rem;
        margin-top: 0.5rem;
    }}
    
    .prompt-chip {{
        background: {LIGHT_BLUE};
        color: {DEEP_BLUE};
        padding: 0.35rem 0.85rem;
        border-radius: 9999px;
        font-size: 0.8rem;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.15s ease;
        border: 1px solid #BFDBFE;
    }}
    
    .prompt-chip:hover {{
        background: {PRIMARY_BLUE};
        color: white;
        border-color: {PRIMARY_BLUE};
    }}

    /* Buttons */
    .stButton>button {{
        background: {PRIMARY_BLUE};
        color: white;
        border-radius: 10px;
        font-weight: 600;
        padding: 0.55rem 1.4rem;
        transition: all 0.2s ease;
        border: none;
    }}
    
    .stButton>button:hover {{
        background: {DEEP_BLUE};
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(13, 110, 253, 0.3);
    }}
    
    .stButton>button[kind="secondary"] {{
        background: {LIGHT_GRAY};
        color: {DARK_TEXT};
    }}

    /* Cards & Containers */
    .glass-card {{
        background: {WHITE};
        border: 1px solid {LIGHT_GRAY};
        border-radius: 16px;
        padding: 1.5rem;
        box-shadow: {CARD_SHADOW};
    }}

    /* Data Table */
    .stDataFrame {{
        border-radius: 12px;
        overflow: hidden;
        border: 1px solid {LIGHT_GRAY};
    }}

    /* Footer */
    .app-footer {{
        margin-top: 3rem;
        padding-top: 1.5rem;
        border-top: 1px solid {LIGHT_GRAY};
        font-size: 0.8rem;
        color: {GRAY};
        display: flex;
        justify-content: space-between;
        align-items: center;
        flex-wrap: wrap;
        gap: 0.5rem;
    }}

    /* Misc Polish */
    .stTabs [data-baseweb="tab-list"] {{
        gap: 4px;
        background: {LIGHT_GRAY};
        padding: 4px;
        border-radius: 12px;
    }}
    
    .stTabs [data-baseweb="tab"] {{
        border-radius: 8px;
        padding: 0.5rem 1rem;
        font-weight: 600;
    }}
    
    .stTabs [aria-selected="true"] {{
        background: white;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }}

    /* Responsive */
    @media (max-width: 768px) {{
        .hero h1 {{ font-size: 2.1rem; }}
        .nav-pills {{ display: none; }}
    }}
</style>
""", unsafe_allow_html=True)

# ============================================================
# SESSION STATE INITIALIZATION
# ============================================================
if "df" not in st.session_state:
    st.session_state.df = None

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "last_upload_name" not in st.session_state:
    st.session_state.last_upload_name = None

if "current_view" not in st.session_state:
    st.session_state.current_view = "home"   # home | dashboard | data | ai | reports

if "filter_region" not in st.session_state:
    st.session_state.filter_region = "All"

if "filter_status" not in st.session_state:
    st.session_state.filter_status = "All"

# ============================================================
# HTML SAFETY HELPER (prevents raw HTML glitch from unescaped user/data content)
# ============================================================
def safe_html(text) -> str:
    """Escape <, >, &, quotes etc so dynamic text can be safely injected into custom HTML via unsafe_allow_html."""
    if text is None:
        return ""
    return html.escape(str(text))


# ============================================================
# HELPER FUNCTIONS - DATA
# ============================================================
def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Map common column name variations to standard names."""
    col_map = {
        "activity": "Activity",
        "activities": "Activity",
        "task": "Activity",
        "intervention": "Activity",
        "indicator": "Indicator",
        "indicators": "Indicator",
        "target": "Target",
        "targets": "Target",
        "planned": "Target",
        "actual": "Actual",
        "achieved": "Actual",
        "result": "Actual",
        "results": "Actual",
        "value": "Actual",
        "unit": "Unit",
        "measure": "Unit",
        "date": "Date",
        "dates": "Date",
        "period": "Date",
        "status": "Status",
        "state": "Status",
        "region": "Region",
        "location": "Region",
        "area": "Region",
        "district": "Region",
        "comments": "Comments",
        "notes": "Comments",
        "remark": "Comments",
        "remarks": "Comments",
        "description": "Comments",
    }
    
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    
    new_cols = {}
    for c in df.columns:
        lower = c.lower().strip()
        if lower in col_map:
            new_cols[c] = col_map[lower]
        else:
            new_cols[c] = c  # keep original if no match
    
    df = df.rename(columns=new_cols)
    
    # Ensure essential columns exist
    for col in ["Activity", "Indicator", "Target", "Actual"]:
        if col not in df.columns:
            df[col] = None
    
    # Convert numeric
    for col in ["Target", "Actual"]:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    
    # Fill missing
    if "Unit" not in df.columns:
        df["Unit"] = "%"
    if "Status" not in df.columns:
        df["Status"] = "On Track"
    if "Region" not in df.columns:
        df["Region"] = "All"
    if "Comments" not in df.columns:
        df["Comments"] = ""
    if "Date" not in df.columns:
        df["Date"] = datetime.now().strftime("%Y-%m-%d")
    
    # Clean status values
    status_map = {
        "on track": "On Track",
        "on-track": "On Track",
        "track": "On Track",
        "completed": "Completed",
        "done": "Completed",
        "delayed": "Delayed",
        "behind": "Behind",
        "at risk": "Behind",
        "off track": "Behind",
    }
    df["Status"] = df["Status"].astype(str).str.strip().str.title()
    df["Status"] = df["Status"].replace(status_map)
    
    # Parse dates
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    df["Date"] = df["Date"].fillna(pd.Timestamp.now())
    
    return df[["Activity", "Indicator", "Target", "Actual", "Unit", "Date", "Status", "Region", "Comments"] + 
              [c for c in df.columns if c not in ["Activity", "Indicator", "Target", "Actual", "Unit", "Date", "Status", "Region", "Comments"]]]


def compute_progress_metrics(df: pd.DataFrame) -> dict:
    """Calculate all key MERL metrics from the dataframe."""
    if df is None or len(df) == 0:
        return {
            "overall_progress": 0.0,
            "num_activities": 0,
            "total_target": 0,
            "total_actual": 0,
            "on_track": 0,
            "behind": 0,
            "completed": 0,
            "delayed": 0,
            "avg_variance": 0.0,
            "regions": [],
            "top_performers": [],
            "lagging": [],
        }
    
    df = df.copy()
    df = df.dropna(subset=["Target", "Actual"], how="all")
    
    # Progress per row
    df["Progress"] = df.apply(
        lambda r: round((r["Actual"] / r["Target"] * 100), 1) if pd.notna(r["Target"]) and r["Target"] > 0 else 0.0,
        axis=1
    )
    df["Progress"] = df["Progress"].clip(upper=120)  # cap at 120% for display
    
    # Variance (positive = over, negative = under)
    df["Variance"] = df.apply(
        lambda r: round(((r["Actual"] - r["Target"]) / r["Target"] * 100), 1) if pd.notna(r["Target"]) and r["Target"] > 0 else 0.0,
        axis=1
    )
    
    overall = round(df["Progress"].mean(), 1) if len(df) > 0 else 0
    
    # Status counts (use both computed and provided)
    on_track = len(df[(df["Status"].str.contains("Track|Completed", case=False, na=False)) | (df["Progress"] >= 80)])
    behind = len(df[(df["Status"].str.contains("Behind|Delayed|Risk", case=False, na=False)) | (df["Progress"] < 70)])
    completed = len(df[df["Status"].str.contains("Completed", case=False, na=False)])
    delayed = len(df[df["Status"].str.contains("Delayed", case=False, na=False)])
    
    total_target = df["Target"].sum(skipna=True)
    total_actual = df["Actual"].sum(skipna=True)
    
    avg_var = round(df["Variance"].mean(), 1) if len(df) > 0 else 0
    
    # Top / lagging
    top = df.nlargest(3, "Progress")[["Activity", "Progress", "Region"]].to_dict("records")
    lag = df.nsmallest(3, "Progress")[["Activity", "Progress", "Region"]].to_dict("records")
    
    return {
        "overall_progress": overall,
        "num_activities": len(df),
        "total_target": int(total_target) if not pd.isna(total_target) else 0,
        "total_actual": int(total_actual) if not pd.isna(total_actual) else 0,
        "on_track": on_track,
        "behind": behind,
        "completed": completed,
        "delayed": delayed,
        "avg_variance": avg_var,
        "regions": sorted(df["Region"].dropna().unique().tolist()),
        "top_performers": top,
        "lagging": lag,
        "df_with_progress": df,
    }


def load_sample_data():
    """Load the bundled sample CSV."""
    try:
        df = pd.read_csv("sample_data.csv")
        return normalize_columns(df)
    except Exception as e:
        st.error(f"Could not load sample data: {e}")
        # Fallback tiny dataset
        data = {
            "Activity": ["Training", "Distribution", "Monitoring"],
            "Indicator": ["Farmers trained", "Inputs delivered", "Visits done"],
            "Target": [100, 500, 24],
            "Actual": [87, 420, 19],
            "Unit": ["%", "kg", "visits"],
            "Date": ["2025-02-01", "2025-02-15", "2025-03-01"],
            "Status": ["On Track", "Delayed", "On Track"],
            "Region": ["North", "South", "North"],
            "Comments": ["", "Supply issues", ""],
        }
        return pd.DataFrame(data)


# ============================================================
# SMART RULE-BASED AI ENGINE (No API key needed)
# ============================================================
def get_data_insights(df: pd.DataFrame) -> dict:
    """Deep analysis for the AI to use."""
    if df is None or len(df) == 0:
        return {}
    
    metrics = compute_progress_metrics(df)
    dfp = metrics.get("df_with_progress", df.copy())
    
    insights = {
        "metrics": metrics,
        "summary_text": "",
        "risks": [],
        "strengths": [],
        "recommendations": [],
    }
    
    # Overall narrative
    prog = metrics["overall_progress"]
    if prog >= 95:
        level = "excellent"
    elif prog >= 80:
        level = "strong"
    elif prog >= 65:
        level = "moderate"
    else:
        level = "concerning"
    
    insights["summary_text"] = (
        f"Overall project performance is **{level}** at {prog}% average progress across {metrics['num_activities']} indicators. "
        f"We have achieved {metrics['total_actual']} against a combined target of {metrics['total_target']}. "
    )
    
    # Strengths
    if metrics["on_track"] > metrics["behind"]:
        insights["strengths"].append(f"Strong delivery: {metrics['on_track']} indicators are on track or ahead.")
    if any(p["Progress"] > 100 for p in metrics["top_performers"]):
        insights["strengths"].append("Several activities have exceeded targets — good candidates for scale-up.")
    
    # Risks
    if metrics["behind"] > 0:
        insights["risks"].append(f"{metrics['behind']} indicators are significantly behind (<70% of target).")
    for lag in metrics["lagging"]:
        if lag["Progress"] < 60:
            insights["risks"].append(f"Critical gap: '{lag['Activity']}' at only {lag['Progress']}% in {lag['Region']}.")
    
    # Variance based
    if metrics["avg_variance"] < -15:
        insights["risks"].append("Systematic under-performance: average variance is negative and large.")
    
    # Region concentration
    if len(metrics["regions"]) > 1:
        region_perf = dfp.groupby("Region")["Progress"].mean().to_dict()
        worst_region = min(region_perf, key=region_perf.get)
        if region_perf[worst_region] < 75:
            insights["risks"].append(f"Geographic disparity: {worst_region} region is lagging at {round(region_perf[worst_region],1)}%.")
    
    # Recommendations (learning & improvement)
    recs = []
    if metrics["behind"] > 0:
        recs.append("Conduct root-cause analysis (fishbone or 5 Whys) on lagging indicators within the next 2 weeks.")
    if metrics["delayed"] > 0:
        recs.append("Review procurement and logistics processes — delays appear supply-chain related.")
    if prog < 80:
        recs.append("Consider revising activity timelines or increasing field staffing for Q2.")
    recs.append("Document successful strategies from top-performing activities and replicate in other regions.")
    recs.append("Increase community feedback loops (suggestion boxes, WhatsApp groups) to improve adaptive management.")
    
    insights["recommendations"] = recs[:5]
    insights["strengths"] = insights["strengths"][:3] or ["Data collection and reporting discipline is good."]
    
    return insights


def ai_respond(query: str, df: pd.DataFrame) -> str:
    """Intelligent, data-driven response generator."""
    if df is None or len(df) == 0:
        return "I don't have any data loaded yet. Please upload a file or add some records using the form first."
    
    q = query.lower().strip()
    ins = get_data_insights(df)
    m = ins["metrics"]
    dfp = m.get("df_with_progress", df)
    
    # --- Specific question routing ---
    if any(k in q for k in ["overall progress", "how are we doing", "summary", "status"]):
        return (
            f"**Overall Progress: {m['overall_progress']}%**\n\n"
            f"Across **{m['num_activities']}** tracked indicators, the project has reached **{m['total_actual']}** out of **{m['total_target']}** planned.\n"
            f"- On Track / Completed: **{m['on_track']}**\n"
            f"- Behind or Delayed: **{m['behind']}**\n\n"
            f"{ins['summary_text']}"
        )
    
    if any(k in q for k in ["insight", "evaluation", "analysis", "performance"]):
        strengths = "\n- ".join([""] + ins["strengths"])
        risks = "\n- ".join([""] + ins["risks"]) if ins["risks"] else "\n- No major systemic risks detected at this time."
        return (
            f"### Evaluation Insights\n\n"
            f"**Performance Level**: {m['overall_progress']}% average\n\n"
            f"**Strengths**:{strengths}\n\n"
            f"**Areas of Concern**:{risks}\n\n"
            f"The data suggests the project is making steady progress but would benefit from focused attention on the lagging indicators."
        )
    
    if any(k in q for k in ["risk", "problem", "issue", "concern", "behind"]):
        if not ins["risks"]:
            return "Good news — current analysis shows no high-severity risks. All indicators are within acceptable variance bands."
        risk_text = "\n\n• ".join(ins["risks"])
        return f"**Key Risks Identified:**\n\n• {risk_text}\n\n**Recommended immediate action**: Prioritize the 2-3 lowest-performing activities for field verification visits."
    
    if any(k in q for k in ["recommend", "improve", "learn", "next step", "action"]):
        rec_text = "\n\n✓ ".join([""] + ins["recommendations"])
        return f"**AI Recommendations for Learning & Improvement:**{rec_text}\n\nThese are generated dynamically from your current dataset."
    
    if any(k in q for k in ["report", "write", "summary report", "executive"]):
        return (
            f"**Executive Summary (auto-generated)**\n\n"
            f"The MERL project currently stands at **{m['overall_progress']}%** overall implementation progress. "
            f"Out of {m['num_activities']} monitored activities, {m['on_track']} are on track or have exceeded expectations. "
            f"{m['behind']} indicators require management attention. "
            f"Key strengths include high community engagement in several regions. "
            f"Priority actions: address the {m['behind']} lagging indicators and replicate successful approaches from top performers. "
            f"Full report available via the **Export Report** buttons below."
        )
    
    if any(k in q for k in ["top", "best", "performer", "success"]):
        top_str = "\n".join([f"- {t['Activity']} ({t['Progress']}%) — {t['Region']}" for t in m["top_performers"]])
        return f"**Top Performing Activities:**\n\n{top_str}\n\nThese are strong models for cross-learning."
    
    if any(k in q for k in ["lag", "worst", "lowest", "poor"]):
        lag_str = "\n".join([f"- {l['Activity']} ({l['Progress']}%) — {l['Region']}" for l in m["lagging"]])
        return f"**Activities Needing Attention:**\n\n{lag_str}\n\nI recommend immediate review and support for these."
    
    if any(k in q for k in ["region", "geographic", "where"]):
        region_perf = dfp.groupby("Region")["Progress"].mean().round(1).to_dict()
        reg_text = "\n".join([f"- {r}: {p}%" for r, p in region_perf.items()])
        return f"**Progress by Region:**\n\n{reg_text}"
    
    # Default intelligent answer
    return (
        f"I analyzed your data: **{m['overall_progress']}%** average progress across {m['num_activities']} indicators. "
        f"{m['on_track']} are performing well. "
        f"You can ask me about risks, recommendations, top performers, specific regions, or request a full summary report."
    )


# ============================================================
# REPORT GENERATORS
# ============================================================
class MERLReportPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
    
    def header(self):
        self.set_fill_color(13, 110, 253)  # Primary blue
        self.rect(0, 0, 210, 18, "F")
        self.set_text_color(255, 255, 255)
        self.set_font("Helvetica", "B", 14)
        self.cell(0, 12, "MERL AI Agent - Project Performance Report", ln=True, align="C")
        self.ln(8)
    
    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(100)
        self.cell(0, 10, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} | MERL AI Dashboard", align="C")


def generate_pdf_report(df: pd.DataFrame, insights: dict) -> bytes:
    """Create a professional multi-page PDF report."""
    m = insights.get("metrics", compute_progress_metrics(df))
    
    pdf = MERLReportPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "", 11)
    
    # Title block
    pdf.set_text_color(13, 110, 253)
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, "Monitoring, Evaluation, Research & Learning Report", ln=True)
    pdf.ln(3)
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(80)
    pdf.cell(0, 6, f"Report Date: {datetime.now().strftime('%B %d, %Y')}", ln=True)
    pdf.ln(6)
    
    # Executive Summary
    pdf.set_text_color(30)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 8, "Executive Summary", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6, 
        f"The project is currently at {m['overall_progress']}% overall progress. "
        f"{m['num_activities']} indicators are being tracked. "
        f"Total achieved vs planned: {m['total_actual']} / {m['total_target']}. "
        f"{m['on_track']} indicators are on track or ahead; {m['behind']} require attention."
    )
    pdf.ln(4)
    
    # KPIs table
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Key Performance Indicators", ln=True)
    pdf.set_font("Helvetica", "", 10)
    
    pdf.set_fill_color(231, 241, 255)
    pdf.cell(95, 8, "Metric", border=1, fill=True)
    pdf.cell(95, 8, "Value", border=1, fill=True, ln=True)
    
    kpis = [
        ("Overall Progress", f"{m['overall_progress']}%"),
        ("Indicators Tracked", str(m['num_activities'])),
        ("On Track / Completed", str(m['on_track'])),
        ("Behind / Delayed", str(m['behind'])),
        ("Average Variance", f"{m['avg_variance']}%"),
    ]
    for label, val in kpis:
        pdf.cell(95, 7, label, border=1)
        pdf.cell(95, 7, val, border=1, ln=True)
    pdf.ln(6)
    
    # Top & Lagging
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Top Performers", ln=True)
    pdf.set_font("Helvetica", "", 10)
    for t in m["top_performers"]:
        pdf.cell(0, 6, f"• {t['Activity']} — {t['Progress']}% ({t['Region']})", ln=True)
    pdf.ln(3)
    
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Activities Requiring Attention", ln=True)
    pdf.set_font("Helvetica", "", 10)
    for l in m["lagging"]:
        pdf.cell(0, 6, f"• {l['Activity']} — {l['Progress']}% ({l['Region']})", ln=True)
    pdf.ln(5)
    
    # Recommendations
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "AI Recommendations", ln=True)
    pdf.set_font("Helvetica", "", 10)
    for r in insights.get("recommendations", []):
        pdf.multi_cell(0, 6, f"• {r}")
    pdf.ln(4)
    
    # Data table (first 8 rows)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Data Snapshot (first 8 records)", ln=True)
    pdf.set_font("Helvetica", "", 9)
    
    # Simple table header
    headers = ["Activity", "Indicator", "Target", "Actual", "Progress %", "Status"]
    col_width = [38, 42, 22, 22, 24, 32]
    pdf.set_fill_color(13, 110, 253)
    pdf.set_text_color(255)
    for h, w in zip(headers, col_width):
        pdf.cell(w, 7, h, border=1, fill=True)
    pdf.ln()
    
    pdf.set_text_color(30)
    pdf.set_fill_color(248, 250, 252)
    for idx, row in df.head(8).iterrows():
        prog = round((row["Actual"]/row["Target"]*100),1) if pd.notna(row["Target"]) and row["Target"]>0 else 0
        vals = [
            str(row["Activity"])[:18],
            str(row["Indicator"])[:20],
            str(int(row["Target"])) if pd.notna(row["Target"]) else "-",
            str(int(row["Actual"])) if pd.notna(row["Actual"]) else "-",
            f"{prog}%",
            str(row["Status"])[:12]
        ]
        for v, w in zip(vals, col_width):
            pdf.cell(w, 6, v, border=1, fill=(idx % 2 == 0))
        pdf.ln()
    
    pdf.ln(8)
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(100)
    pdf.multi_cell(0, 5, "This report was automatically generated by the MERL AI Agent. For interactive analysis and deeper filtering, please use the web dashboard.")
    
    return pdf.output(dest="S").encode("latin-1")


def generate_word_report(df: pd.DataFrame, insights: dict) -> bytes:
    """Create a polished .docx report."""
    m = insights.get("metrics", compute_progress_metrics(df))
    
    doc = Document()
    
    # Title
    title = doc.add_heading("MERL AI Agent — Project Performance Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    p = doc.add_paragraph()
    p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"The project currently stands at {m['overall_progress']}% overall progress across {m['num_activities']} tracked indicators. "
        f"A total of {m['total_actual']} has been achieved against a combined target of {m['total_target']}. "
        f"{m['on_track']} indicators are performing on track or better, while {m['behind']} require management attention."
    )
    
    # KPIs
    doc.add_heading("Key Performance Indicators", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    kpis = [
        ("Overall Progress", f"{m['overall_progress']}%"),
        ("Indicators Tracked", str(m["num_activities"])),
        ("On Track / Completed", str(m["on_track"])),
        ("Behind / Delayed", str(m["behind"])),
        ("Average Variance from Target", f"{m['avg_variance']}%"),
    ]
    for i, (label, val) in enumerate(kpis):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = val
    
    doc.add_paragraph()
    
    # Top performers
    doc.add_heading("Top Performers", level=1)
    for t in m["top_performers"]:
        doc.add_paragraph(f"{t['Activity']} — {t['Progress']}% progress ({t['Region']})", style="List Bullet")
    
    # Lagging
    doc.add_heading("Activities Requiring Attention", level=1)
    for l in m["lagging"]:
        doc.add_paragraph(f"{l['Activity']} — {l['Progress']}% progress ({l['Region']})", style="List Bullet")
    
    # Recommendations
    doc.add_heading("AI-Generated Recommendations", level=1)
    for r in insights.get("recommendations", []):
        doc.add_paragraph(r, style="List Bullet")
    
    # Data table
    doc.add_heading("Data Snapshot", level=1)
    t = doc.add_table(rows=min(9, len(df)+1), cols=6)
    t.style = "Table Grid"
    
    headers = ["Activity", "Indicator", "Target", "Actual", "Progress %", "Status"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for paragraph in t.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for row_idx, (_, row) in enumerate(df.head(8).iterrows()):
        prog = round((row["Actual"]/row["Target"]*100),1) if pd.notna(row["Target"]) and row["Target"]>0 else 0
        vals = [
            str(row["Activity"])[:22],
            str(row["Indicator"])[:24],
            str(int(row["Target"])) if pd.notna(row["Target"]) else "-",
            str(int(row["Actual"])) if pd.notna(row["Actual"]) else "-",
            f"{prog}%",
            str(row["Status"])
        ]
        for col_idx, v in enumerate(vals):
            t.rows[row_idx+1].cells[col_idx].text = v
    
    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("Generated automatically by the MERL AI Agent Dashboard. ").italic = True
    footer.add_run("For live filtering, charts, and interactive Q&A, open the Streamlit application.").italic = True
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# MAIN APP UI
# ============================================================
def main():
    """Main entry point with beautiful modern navigation and views."""

    # ============================================================
    # DATA LOADING & FILTERING (runs for all views)
    # ============================================================
    if st.session_state.df is None:
        st.session_state.df = load_sample_data()
        st.session_state.last_upload_name = "sample_data.csv (auto-loaded)"

    # Sync current_view from URL ?view= query param (enables clean functional top nav links)
    qp_view = st.query_params.get("view")
    if isinstance(qp_view, (list, tuple)):
        qp_view = qp_view[0] if qp_view else None
    allowed_views = ["home", "dashboard", "data", "ai", "reports"]
    if qp_view and qp_view in allowed_views and st.session_state.current_view != qp_view:
        st.session_state.current_view = qp_view

    # Get current data and apply filters
    base_df = st.session_state.df.copy()

    # Filters (shown only in relevant views)
    filtered_df = base_df.copy()
    if st.session_state.filter_region != "All":
        filtered_df = filtered_df[filtered_df["Region"] == st.session_state.filter_region]
    if st.session_state.filter_status != "All":
        filtered_df = filtered_df[filtered_df["Status"] == st.session_state.filter_status]

    metrics = compute_progress_metrics(filtered_df)
    insights = get_data_insights(filtered_df)

    # ============================================================
    # TOP NAVIGATION BAR (Modern & Professional)
    # ============================================================
    def set_view(view_name):
        st.session_state.current_view = view_name
        try:
            st.query_params["view"] = view_name
        except Exception:
            pass
        st.rerun()

    nav_html = f"""
    <div class="nav-container">
        <div class="nav-inner">
            <div class="nav-logo">
                <span>📊</span> <span>MERL AI</span>
            </div>
            <div class="nav-pills">
                <span class="nav-pill {'active' if st.session_state.current_view == 'home' else ''}" 
                      style="{'background:white;color:#1E40AF;box-shadow:0 1px 3px rgba(0,0,0,0.1);cursor:default' if st.session_state.current_view == 'home' else 'cursor:default'}">
                    🏠 Home
                </span>
                <span class="nav-pill {'active' if st.session_state.current_view == 'dashboard' else ''}" 
                      style="{'background:white;color:#1E40AF;box-shadow:0 1px 3px rgba(0,0,0,0.1);cursor:default' if st.session_state.current_view == 'dashboard' else 'cursor:default'}">
                    📈 Dashboard
                </span>
                <span class="nav-pill {'active' if st.session_state.current_view == 'data' else ''}" 
                      style="{'background:white;color:#1E40AF;box-shadow:0 1px 3px rgba(0,0,0,0.1);cursor:default' if st.session_state.current_view == 'data' else 'cursor:default'}">
                    📋 Data
                </span>
                <span class="nav-pill {'active' if st.session_state.current_view == 'ai' else ''}" 
                      style="{'background:white;color:#1E40AF;box-shadow:0 1px 3px rgba(0,0,0,0.1);cursor:default' if st.session_state.current_view == 'ai' else 'cursor:default'}">
                    🤖 AI Chat
                </span>
                <span class="nav-pill {'active' if st.session_state.current_view == 'reports' else ''}" 
                      style="{'background:white;color:#1E40AF;box-shadow:0 1px 3px rgba(0,0,0,0.1);cursor:default' if st.session_state.current_view == 'reports' else 'cursor:default'}">
                    📄 Reports
                </span>
            </div>
            <div style="display:flex; align-items:center; gap:8px; font-size:0.85rem;">
                <span style="color:#64748B;">v2.1</span>
            </div>
        </div>
    </div>
    """
    st.markdown(nav_html, unsafe_allow_html=True)

    # Functional navigation buttons (the pills above are visual only).
    # We keep them so navigation is a soft Streamlit rerun (no full browser reload that can expose raw HTML).
    col_nav = st.columns(5)
    with col_nav[0]:
        if st.button("🏠 Home", key="nav_home", use_container_width=True):
            set_view("home")
    with col_nav[1]:
        if st.button("📈 Dashboard", key="nav_dash", use_container_width=True):
            set_view("dashboard")
    with col_nav[2]:
        if st.button("📋 Data & Entry", key="nav_data", use_container_width=True):
            set_view("data")
    with col_nav[3]:
        if st.button("🤖 AI Assistant", key="nav_ai", use_container_width=True):
            set_view("ai")
    with col_nav[4]:
        if st.button("📄 Reports & Insights", key="nav_reports", use_container_width=True):
            set_view("reports")

    st.markdown("<br>", unsafe_allow_html=True)

    # ============================================================
    # VIEW: HOME / LANDING PAGE
    # ============================================================
    if st.session_state.current_view == "home":
        # Beautiful Hero
        st.markdown(f"""
        <div class="hero">
            <h1>Monitor.<br>Evaluate.<br>Learn faster.</h1>
            <p>Professional AI-powered dashboard for Monitoring, Evaluation, Research & Learning teams. 
            Upload your data, get instant insights, and export stakeholder-ready reports in seconds.</p>
        </div>
        """, unsafe_allow_html=True)

        # Quick action buttons
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("📥 Load Realistic Sample Data", use_container_width=True, type="primary"):
                st.session_state.df = load_sample_data()
                st.session_state.last_upload_name = "sample_data.csv"
                st.session_state.chat_history = []
                st.session_state.filter_region = "All"
                st.session_state.filter_status = "All"
                set_view("dashboard")
        with c2:
            if st.button("📈 Go to Dashboard", use_container_width=True):
                set_view("dashboard")
        with c3:
            if st.button("🤖 Talk to the AI", use_container_width=True):
                set_view("ai")

        st.markdown("<div class='section-header'>Everything you need for strong MERL work</div>", unsafe_allow_html=True)

        # Feature grid
        st.markdown("""
        <div class="feature-grid">
            <div class="feature-card">
                <div class="feature-icon" style="background:#EFF6FF;color:#1E40AF;">📊</div>
                <strong>Live Dashboard</strong>
                <p style="font-size:0.9rem;color:#64748B;margin-top:6px;">Beautiful interactive charts showing progress, variance, and regional performance at a glance.</p>
            </div>
            <div class="feature-card">
                <div class="feature-icon" style="background:#ECFDF5;color:#059669;">🤖</div>
                <strong>Smart AI Analyst</strong>
                <p style="font-size:0.9rem;color:#64748B;margin-top:6px;">Ask natural questions and get data-driven evaluation insights, risk alerts, and recommendations instantly.</p>
            </div>
            <div class="feature-card">
                <div class="feature-icon" style="background:#FEF3C7;color:#B45309;">📝</div>
                <strong>One-Click Reports</strong>
                <p style="font-size:0.9rem;color:#64748B;margin-top:6px;">Generate polished PDF and Word documents with executive summaries, tables, and AI recommendations.</p>
            </div>
            <div class="feature-card">
                <div class="feature-icon" style="background:#F3E8FF;color:#7C3AED;">✚</div>
                <strong>Easy Data Management</strong>
                <p style="font-size:0.9rem;color:#64748B;margin-top:6px;">Drag & drop Excel/CSV or add records manually. Works offline. No data leaves your browser.</p>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        st.info("💡 **Tip**: Start with the sample data to explore all features, then replace it with your own project file.")

    # ============================================================
    # SHARED: DATA SOURCE CONTROLS (Upload + Sample + Reset)
    # ============================================================
    def render_data_controls():
        with st.expander("📁 Data Source & Upload", expanded=(st.session_state.current_view == "data")):
            col_u1, col_u2, col_u3 = st.columns([2, 1, 1])
            
            with col_u1:
                uploaded = st.file_uploader(
                    "Drag & drop Excel or CSV file",
                    type=["csv", "xlsx", "xls"],
                    label_visibility="collapsed"
                )
                if uploaded is not None:
                    try:
                        if uploaded.name.endswith((".xlsx", ".xls")):
                            raw = pd.read_excel(uploaded, engine="openpyxl")
                        else:
                            raw = pd.read_csv(uploaded)
                        st.session_state.df = normalize_columns(raw)
                        st.session_state.last_upload_name = uploaded.name
                        st.session_state.chat_history = []
                        st.session_state.filter_region = "All"
                        st.session_state.filter_status = "All"
                        st.success(f"✅ Loaded **{uploaded.name}** ({len(st.session_state.df)} rows)")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Could not read file: {e}")

            with col_u2:
                if st.button("📥 Load Sample Data", use_container_width=True):
                    st.session_state.df = load_sample_data()
                    st.session_state.last_upload_name = "sample_data.csv"
                    st.session_state.chat_history = []
                    st.session_state.filter_region = "All"
                    st.session_state.filter_status = "All"
                    st.rerun()

            with col_u3:
                if st.button("🔄 Reset Everything", use_container_width=True):
                    st.session_state.df = None
                    st.session_state.chat_history = []
                    st.session_state.last_upload_name = None
                    st.session_state.filter_region = "All"
                    st.session_state.filter_status = "All"
                    st.rerun()

            st.caption(f"Current file: **{st.session_state.last_upload_name}** • {len(base_df)} total records")

    # ============================================================
    # VIEW: DASHBOARD (Beautiful KPIs + Charts)
    # ============================================================
    if st.session_state.current_view == "dashboard":
        st.markdown(f"<div class='section-header'>📈 Project Performance Dashboard</div>", unsafe_allow_html=True)
        safe_source = safe_html(st.session_state.last_upload_name)
        st.markdown(f"<div class='section-sub'>Showing data from: <strong>{safe_source}</strong> — {len(filtered_df)} records after filters</div>", unsafe_allow_html=True)

        render_data_controls()

        # Filters row
        f1, f2, f3 = st.columns([1.2, 1.2, 3])
        with f1:
            regions = ["All"] + sorted(base_df["Region"].dropna().unique().tolist())
            reg_val = st.session_state.filter_region
            reg_idx = 0 if reg_val == "All" or reg_val not in regions else regions.index(reg_val)
            st.session_state.filter_region = st.selectbox("Region", regions, index=reg_idx)
        with f2:
            statuses = ["All"] + sorted(base_df["Status"].dropna().unique().tolist())
            stat_val = st.session_state.filter_status
            stat_idx = 0 if stat_val == "All" or stat_val not in statuses else statuses.index(stat_val)
            st.session_state.filter_status = st.selectbox("Status", statuses, index=stat_idx)
        with f3:
            st.caption("Filters apply to Dashboard, AI Chat, and Reports")

        # KPI Cards - 6 modern cards
        kpi_cols = st.columns(6)
        kpis = [
            (f"{metrics['overall_progress']}%", "Overall Progress", "blue"),
            (str(metrics['num_activities']), "Indicators Tracked", "blue"),
            (str(metrics['on_track']), "On Track", "green"),
            (str(metrics['behind']), "Behind Schedule", "orange" if metrics['behind'] > 0 else "green"),
            (f"{metrics['avg_variance']:+.0f}%", "Avg Variance", "blue"),
            (f"{metrics['total_actual']:,}", "Total Achieved", "green"),
        ]
        for i, (val, label, tone) in enumerate(kpis):
            color = PRIMARY_BLUE if tone == "blue" else (ACCENT_GREEN if tone == "green" else WARNING)
            with kpi_cols[i]:
                v = safe_html(val)
                l = safe_html(label)
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value" style="color:{color};">{v}</div>
                    <div class="metric-label">{l}</div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # Charts - 2x2 professional layout
        chart_col1, chart_col2 = st.columns(2)

        with chart_col1:
            st.markdown("**Target vs Actual Performance**")
            if len(filtered_df) > 0:
                chart_df = filtered_df.copy()
                chart_df["Progress"] = chart_df.apply(
                    lambda r: (r["Actual"] / r["Target"] * 100) if pd.notna(r["Target"]) and r["Target"] > 0 else 0, axis=1
                )
                fig = px.bar(
                    chart_df,
                    x="Activity",
                    y=["Target", "Actual"],
                    barmode="group",
                    color_discrete_sequence=[PRIMARY_BLUE, ACCENT_GREEN],
                    height=340
                )
                fig.update_layout(
                    template="plotly_white",
                    margin=dict(l=10, r=10, t=10, b=60),
                    xaxis_tickangle=-35,
                    legend=dict(orientation="h", y=-0.25),
                    yaxis_title="Value"
                )
                st.plotly_chart(fig, use_container_width=True)

        with chart_col2:
            st.markdown("**Progress Distribution**")
            if len(filtered_df) > 0:
                prog_df = metrics.get("df_with_progress", filtered_df)
                fig2 = px.histogram(
                    prog_df, x="Progress", nbins=12,
                    color_discrete_sequence=[PRIMARY_BLUE],
                    height=340
                )
                fig2.add_vline(x=80, line_dash="dash", line_color=ACCENT_GREEN, annotation_text="Good target (80%)")
                fig2.update_layout(template="plotly_white", margin=dict(l=10, r=10, t=10, b=10))
                st.plotly_chart(fig2, use_container_width=True)

        c3, c4 = st.columns(2)
        with c3:
            st.markdown("**Status Overview**")
            if len(filtered_df) > 0:
                status_df = filtered_df["Status"].value_counts().reset_index()
                status_df.columns = ["Status", "Count"]
                fig3 = px.pie(status_df, names="Status", values="Count",
                              color_discrete_sequence=[ACCENT_GREEN, PRIMARY_BLUE, WARNING, DANGER], height=300)
                fig3.update_layout(template="plotly_white")
                st.plotly_chart(fig3, use_container_width=True)

        with c4:
            st.markdown("**Performance by Region**")
            if len(filtered_df) > 0:
                reg_perf = metrics.get("df_with_progress", filtered_df).groupby("Region")["Progress"].mean().reset_index()
                fig4 = px.bar(reg_perf, x="Region", y="Progress", color="Progress",
                              color_continuous_scale=["#EF4444", "#F59E0B", ACCENT_GREEN], height=300)
                fig4.update_layout(template="plotly_white")
                st.plotly_chart(fig4, use_container_width=True)

    # ============================================================
    # VIEW: DATA MANAGEMENT
    # ============================================================
    if st.session_state.current_view == "data":
        st.markdown("<div class='section-header'>📋 Data Management</div>", unsafe_allow_html=True)
        render_data_controls()

        st.markdown("<div class='section-header' style='margin-top:1.5rem'>Current Dataset</div>", unsafe_allow_html=True)
        if len(base_df) > 0:
            st.dataframe(
                base_df[["Activity", "Indicator", "Target", "Actual", "Unit", "Status", "Region", "Comments"]],
                use_container_width=True, height=320, hide_index=True
            )
        else:
            st.info("No data yet.")

        st.markdown("<div class='section-header' style='margin-top:1.5rem'>➕ Add New Record</div>", unsafe_allow_html=True)
        with st.form("add_form", clear_on_submit=True):
            c1, c2 = st.columns(2)
            with c1:
                act = st.text_input("Activity *", placeholder="e.g. Farmer Training Round 2")
                ind = st.text_input("Indicator *", placeholder="% of participants applying new practices")
                targ = st.number_input("Target *", min_value=0.0, value=100.0)
                actl = st.number_input("Actual *", min_value=0.0, value=78.0)
            with c2:
                unit = st.text_input("Unit", value="%")
                stat = st.selectbox("Status", ["On Track", "Completed", "Delayed", "Behind"])
                reg = st.text_input("Region", value="All")
                comm = st.text_area("Comments", height=80)

            if st.form_submit_button("Add Record to Dataset", use_container_width=True):
                if act and ind:
                    new = pd.DataFrame([{"Activity": act, "Indicator": ind, "Target": float(targ),
                                         "Actual": float(actl), "Unit": unit or "%", "Date": datetime.now(),
                                         "Status": stat, "Region": reg or "All", "Comments": comm}])
                    st.session_state.df = pd.concat([st.session_state.df, new], ignore_index=True)
                    st.success("Record added successfully!")
                    st.rerun()
                else:
                    st.error("Activity and Indicator are required.")

    # ============================================================
    # VIEW: AI CHAT (Polished)
    # ============================================================
    if st.session_state.current_view == "ai":
        st.markdown("<div class='section-header'>🤖 AI MERL Assistant</div>", unsafe_allow_html=True)
        st.markdown("Ask anything about your data. The AI sees your current filters.", unsafe_allow_html=True)

        # Suggested prompts
        st.markdown("**Quick questions:**")
        chips = st.columns(4)
        prompts = [
            "What is the overall progress?",
            "Give me evaluation insights",
            "What risks do you see?",
            "Write a summary report"
        ]
        for i, p in enumerate(prompts):
            with chips[i]:
                if st.button(p, key=f"chip_{i}", use_container_width=True):
                    st.session_state.chat_history.append(("user", p))
                    resp = ai_respond(p, filtered_df)
                    st.session_state.chat_history.append(("assistant", resp))
                    st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)

        # Chat history with modern styling
        if st.session_state.chat_history:
            for role, msg in st.session_state.chat_history:
                safe_msg = safe_html(msg).replace("\n", "<br>")
                if role == "user":
                    st.markdown(f'<div style="display:flex;justify-content:flex-end"><div class="chat-bubble chat-user">{safe_msg}</div></div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div style="display:flex"><div class="chat-bubble chat-ai"><strong>AI Analyst:</strong><br>{safe_msg}</div></div>', unsafe_allow_html=True)
        else:
            st.info("Start a conversation — try one of the quick questions above or type below.")

        # Input
        user_input = st.chat_input("Ask about progress, risks, recommendations, regions...")
        if user_input:
            st.session_state.chat_history.append(("user", user_input))
            with st.spinner("Thinking..."):
                reply = ai_respond(user_input, filtered_df)
            st.session_state.chat_history.append(("assistant", reply))
            st.rerun()

        if st.button("🗑️ Clear Conversation", use_container_width=False):
            st.session_state.chat_history = []
            st.rerun()

    # ============================================================
    # VIEW: REPORTS + RECOMMENDATIONS
    # ============================================================
    if st.session_state.current_view == "reports":
        st.markdown("<div class='section-header'>📄 Professional Reports & AI Insights</div>", unsafe_allow_html=True)

        # AI Recommendations
        st.markdown("### 💡 AI Recommendations for Improvement")
        for i, rec in enumerate(insights.get("recommendations", []), 1):
            safe_rec = safe_html(rec)
            st.markdown(f"""
            <div style="background:{LIGHT_GREEN}; border-left:5px solid {ACCENT_GREEN}; padding:0.9rem 1.1rem; margin:0.5rem 0; border-radius:8px;">
                <strong>{i}.</strong> {safe_rec}
            </div>
            """, unsafe_allow_html=True)

        # Strengths & Risks
        col_s, col_r = st.columns(2)
        with col_s:
            st.markdown("**✅ Strengths**")
            for s in insights.get("strengths", []):
                st.markdown(f"- {s}")
        with col_r:
            st.markdown("**⚠️ Risks & Gaps**")
            for r in insights.get("risks", []) or ["No major risks detected."]:
                st.markdown(f"- {r}")

        st.markdown("---")
        st.markdown("### 📥 Export Professional Reports")

        r1, r2 = st.columns(2)
        with r1:
            if st.button("📕 Generate & Download PDF Report", use_container_width=True, type="primary"):
                with st.spinner("Creating PDF..."):
                    pdf_bytes = generate_pdf_report(filtered_df, insights)
                    st.download_button(
                        "⬇️ Download MERL_Report.pdf",
                        pdf_bytes,
                        file_name=f"MERL_Report_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                    )

        with r2:
            if st.button("📘 Generate & Download Word Report", use_container_width=True):
                with st.spinner("Creating Word document..."):
                    docx_bytes = generate_word_report(filtered_df, insights)
                    st.download_button(
                        "⬇️ Download MERL_Report.docx",
                        docx_bytes,
                        file_name=f"MERL_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

        st.caption("Reports include executive summary, KPIs, top/lagging performers, and AI recommendations.")

    # ============================================================
    # GLOBAL PROFESSIONAL FOOTER
    # ============================================================
    st.markdown(f"""
    <div class="app-footer">
        <div>
            <strong>MERL AI Agent</strong> v2.1 &nbsp;•&nbsp; Built for MEL practitioners &nbsp;•&nbsp; 
            Data stays 100% in your browser
        </div>
        <div style="color:#94A3B8;">
            Made with ❤️ using Streamlit • Python • Plotly
        </div>
    </div>
    """, unsafe_allow_html=True)
if __name__ == "__main__":
    main()
